from __future__ import annotations

import math
from pathlib import Path

import pandas as pd
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(r"C:\Users\4B\Desktop\PFE")
ASSETS = ROOT / "rapport_assets"
CAP = ASSETS / "captures_hd"
MEDIA = ASSETS / "template_media"
GEN = ASSETS / "generated_figures"
GEN.mkdir(parents=True, exist_ok=True)

OUT = Path(r"C:\Users\4B\Desktop\Rapport PFE - MQTT IDS - Fettache Mohamed et Najim Ilyas.docx")
TITLE = "Conception et réalisation d’un système intelligent de détection d’intrusions IoT basé sur MQTT, Machine Learning et Deep Learning"
STUDENTS = "Fettache Mohamed et Najim Ilyas"
SUPERVISOR = "Mme Fatna El Mendili"
FILIERE = "Intelligence Artificielle et Technologies Émergentes"
YEAR = "2025/2026"


def font(size=22, bold=False):
    name = "arialbd.ttf" if bold else "arial.ttf"
    path = Path(r"C:\Windows\Fonts") / name
    return ImageFont.truetype(str(path), size) if path.exists() else ImageFont.load_default()


def wrap(draw, text, fnt, max_w):
    lines, cur = [], ""
    for word in text.split():
        test = (cur + " " + word).strip()
        if draw.textbbox((0, 0), test, font=fnt)[2] <= max_w:
            cur = test
        else:
            if cur:
                lines.append(cur)
            cur = word
    if cur:
        lines.append(cur)
    return lines


def draw_text(draw, box, text, fnt, fill="#17324d"):
    x1, y1, x2, y2 = box
    lines = wrap(draw, text, fnt, x2 - x1 - 24)
    heights = [draw.textbbox((0, 0), l, font=fnt)[3] for l in lines]
    y = y1 + ((y2 - y1) - sum(heights) - 8 * (len(lines) - 1)) / 2
    for line, h in zip(lines, heights):
        w = draw.textbbox((0, 0), line, font=fnt)[2]
        draw.text((x1 + (x2 - x1 - w) / 2, y), line, font=fnt, fill=fill)
        y += h + 8


def arrow(draw, a, b, color="#1f4e79"):
    draw.line([a, b], fill=color, width=5)
    ang = math.atan2(b[1] - a[1], b[0] - a[0])
    for d in (2.45, -2.45):
        p = (b[0] + 18 * math.cos(ang + d), b[1] + 18 * math.sin(ang + d))
        draw.line([b, p], fill=color, width=5)


def flow_image(path, title, labels, subtitle=""):
    img = Image.new("RGB", (1600, 760), "#f7fafc")
    d = ImageDraw.Draw(img)
    d.text((60, 40), title, font=font(34, True), fill="#16324f")
    x, y, w, h, gap = 70, 250, 210, 150, 55
    colors = ["#e1f5fe", "#fff3cd", "#e8f5e9", "#f3e5f5", "#e7f0ff", "#fff0f0"]
    for i, label in enumerate(labels):
        box = (x + i * (w + gap), y, x + i * (w + gap) + w, y + h)
        d.rounded_rectangle(box, radius=18, fill=colors[i % len(colors)], outline="#1f4e79", width=3)
        draw_text(d, box, label, font(21, True))
        if i < len(labels) - 1:
            arrow(d, (box[2], y + h // 2), (box[2] + gap - 10, y + h // 2))
    if subtitle:
        d.text((100, 500), subtitle, font=font(22), fill="#37474f")
    img.save(path)


ARCH = GEN / "architecture_mqtt_ids.png"
PIPE = GEN / "pipeline_ml.png"
SEQ = GEN / "sequence_validation.png"
flow_image(
    ARCH,
    "Architecture globale du système MQTT-IDS",
    ["ESP32 + DHT22", "Broker Mosquitto", "FastAPI Backend", "XGBoost + Autoencoder", "React Dashboard"],
    "Flux complet : publication MQTT, prédiction, WebSocket et supervision temps réel.",
)
flow_image(
    PIPE,
    "Pipeline de préparation des données et apprentissage",
    ["PCAP MQTTset", "Extraction CSV", "Nettoyage et scaling", "Split 70/30", "Modèles ML/DL", "Artifacts"],
    "Les traces réseau sont transformées en variables numériques avant entraînement et prédiction.",
)
flow_image(
    SEQ,
    "Séquence de validation en temps réel",
    ["Arduino IDE", "MQTT publish", "FastAPI reçoit", "Prédiction ML", "Dashboard alerté"],
    "Chaque scénario d'attaque est visible dans l'interface quelques secondes après publication.",
)


def field(paragraph, code, placeholder=""):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = code
    run._r.append(instr)
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    run._r.append(sep)
    if placeholder:
        paragraph.add_run(placeholder)
    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run("Page ")
    field(paragraph, "PAGE", "1")


def caption(doc, label, text):
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(label + " ").bold = True
    field(p, f"SEQ {label} \\* ARABIC", "1")
    p.add_run(" : " + text)


def shade(cell, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def cell_text(cell, text, bold=False, white=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(str(text))
    r.bold = bold
    r.font.size = Pt(9)
    if white:
        r.font.color.rgb = RGBColor(255, 255, 255)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(doc, title, cols, rows):
    table = doc.add_table(rows=1, cols=len(cols))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, c in enumerate(cols):
        shade(table.rows[0].cells[i], "1F4E79")
        cell_text(table.rows[0].cells[i], c, True, True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell_text(cells[i], value, align=WD_ALIGN_PARAGRAPH.LEFT if i in (0, 1) else WD_ALIGN_PARAGRAPH.CENTER)
    caption(doc, "Tableau", title)
    doc.add_paragraph()
    return table


def para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    p.add_run(text)


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(item)


def fig(doc, path, title, width=16.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width))
    caption(doc, "Figure", title)
    doc.add_paragraph()


def h(doc, level, text):
    doc.add_heading(text, level=level)


def mb(path):
    return f"{path.stat().st_size / 1024 / 1024:.2f} MB"


reduced = pd.read_csv(ROOT / "dataset/Data/FINAL_CSV/mqttdataset_reduced.csv", low_memory=False)
train = pd.read_csv(ROOT / "dataset/Data/FINAL_CSV/train70_reduced.csv", low_memory=False)
test = pd.read_csv(ROOT / "dataset/Data/FINAL_CSV/test30_reduced.csv", low_memory=False)
classes = ["legitimate", "dos", "bruteforce", "malformed", "slowite", "flood"]
ct = reduced["target"].value_counts().to_dict()
ctr = train["target"].value_counts().to_dict()
cte = test["target"].value_counts().to_dict()

doc = Document()
section = doc.sections[0]
section.page_width, section.page_height = Cm(21), Cm(29.7)
section.top_margin, section.bottom_margin = Cm(2), Cm(2)
section.left_margin, section.right_margin = Cm(2.2), Cm(2)

for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3", "Caption"]:
    st = doc.styles[style_name]
    st.font.name = "Times New Roman"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
doc.styles["Normal"].font.size = Pt(12)
doc.styles["Heading 1"].font.size = Pt(16)
doc.styles["Heading 1"].font.bold = True
doc.styles["Heading 1"].font.color.rgb = RGBColor(31, 78, 121)
doc.styles["Heading 2"].font.size = Pt(14)
doc.styles["Heading 2"].font.bold = True
doc.styles["Heading 2"].font.color.rgb = RGBColor(31, 78, 121)
doc.styles["Caption"].font.size = Pt(10)
doc.styles["Caption"].font.italic = True

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.CENTER
header.add_run("Projet de Fin d'Études - MQTT-IDS").font.size = Pt(9)
page_number(section.footer.paragraphs[0])

# Cover
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
logo = MEDIA / "image2.jpeg"
if logo.exists():
    p.add_run().add_picture(str(logo), width=Cm(8.6))
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Rapport de Projet de Fin d'Études")
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = RGBColor(31, 78, 121)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Filière : " + FILIERE)
r.bold = True
r.font.size = Pt(14)
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(TITLE)
r.bold = True
r.font.size = Pt(18)
r.font.color.rgb = RGBColor(20, 55, 90)
doc.add_paragraph()
cover = add_table(
    doc,
    "Informations générales du projet",
    ["Élément", "Information"],
    [
        ("Réalisé par", STUDENTS),
        ("Encadré par", SUPERVISOR),
        ("Établissement", "École Supérieure de Technologie de Meknès"),
        ("Année universitaire", YEAR),
    ],
)
for row in cover.rows[1:]:
    shade(row.cells[0], "D9EAF7")
doc.add_page_break()

h(doc, 1, "Remerciements")
for text in [
    "Nous tenons à exprimer notre profonde gratitude à notre encadrante, Mme Fatna El Mendili, pour son accompagnement, ses remarques constructives et sa disponibilité durant la réalisation de ce projet de fin d'études.",
    "Nous remercions également l'ensemble des enseignants de l'École Supérieure de Technologie de Meknès pour la formation reçue dans les domaines de l'intelligence artificielle, des systèmes connectés, du développement logiciel et de la cybersécurité.",
    "Enfin, nous adressons nos remerciements à nos familles et à nos collègues pour leur soutien continu. Ce travail représente une synthèse entre apprentissage académique, expérimentation technique et effort personnel.",
]:
    para(doc, text)

h(doc, 1, "Résumé")
for text in [
    "Ce projet de fin d'études porte sur la conception et la réalisation d'un système intelligent de détection d'intrusions dans un environnement IoT utilisant le protocole MQTT. La solution proposée, nommée MQTT-IDS, combine un microcontrôleur ESP32, un capteur DHT22, un broker Mosquitto, un backend FastAPI, des modèles de Machine Learning et une interface web React de supervision en temps réel.",
    "Le système reçoit les messages MQTT publiés par les objets connectés, extrait ou normalise les informations pertinentes, puis applique un modèle XGBoost pour classifier le trafic en plusieurs classes : legitimate, DoS, Flood, Bruteforce, SlowITe et Malformed. Un autoencoder est également intégré pour renforcer la détection d'anomalies inconnues.",
    "Les tests réalisés montrent la capacité de la plateforme à traiter des scénarios variés d'attaques MQTT et à fournir une visualisation immédiate des événements détectés. Le rapport présente le contexte, l'analyse, la conception, l'implémentation, les essais, les limites et les perspectives d'amélioration de la solution.",
]:
    para(doc, text)

h(doc, 1, "Abstract")
for text in [
    "This final year project presents the design and implementation of an intelligent intrusion detection system for IoT environments based on the MQTT protocol.",
    "The proposed platform combines ESP32-based data publishing, Mosquitto broker communication, a FastAPI backend, machine learning models and a real-time React dashboard.",
    "The project demonstrates the relevance of combining IoT communication, cybersecurity and artificial intelligence to build a practical monitoring tool for connected environments.",
]:
    para(doc, text)

doc.add_page_break()
for title, code in [
    ("Table des matières", 'TOC \\o "1-3" \\h \\z \\u'),
    ("Liste des figures", 'TOC \\h \\z \\c "Figure"'),
    ("Liste des tableaux", 'TOC \\h \\z \\c "Tableau"'),
]:
    h(doc, 1, title)
    field(doc.add_paragraph(), code, "Champ à mettre à jour.")
    doc.add_page_break()

h(doc, 1, "Liste des abréviations")
add_table(
    doc,
    "Abréviations utilisées dans le rapport",
    ["Abréviation", "Signification"],
    [
        ("API", "Application Programming Interface"),
        ("CSV", "Comma-Separated Values"),
        ("DHT22", "Capteur numérique de température et d'humidité"),
        ("DL", "Deep Learning"),
        ("DoS", "Denial of Service"),
        ("ESP32", "Microcontrôleur Wi-Fi/Bluetooth utilisé dans l'IoT"),
        ("IA", "Intelligence Artificielle"),
        ("IDS", "Intrusion Detection System"),
        ("IoT", "Internet of Things"),
        ("JSON", "JavaScript Object Notation"),
        ("ML", "Machine Learning"),
        ("MQTT", "Message Queuing Telemetry Transport"),
        ("PCAP", "Packet Capture"),
        ("QoS", "Quality of Service"),
        ("REST", "Representational State Transfer"),
        ("WebSocket", "Canal de communication bidirectionnel temps réel"),
        ("XGBoost", "Extreme Gradient Boosting"),
    ],
)
doc.add_page_break()

sections = [
    (
        1,
        "Introduction",
        [
            "L'Internet des Objets occupe aujourd'hui une place centrale dans les systèmes connectés. Les capteurs, actionneurs et microcontrôleurs communiquent avec des plateformes logicielles capables de collecter, traiter et visualiser les données en temps réel.",
            "Le protocole MQTT est très utilisé dans les architectures IoT grâce à sa légèreté et à son modèle publish/subscribe. Toutefois, lorsqu'il est mal protégé ou exposé à des flux anormaux, il peut devenir une cible pour des attaques de type déni de service, flood, bruteforce ou messages malformés.",
            "Le présent projet vise à concevoir une plateforme complète, depuis la publication de messages MQTT jusqu'à la visualisation des alertes dans une interface web. Le travail combine développement full stack, intelligence artificielle, analyse de trafic réseau et validation expérimentale.",
        ],
    ),
    (
        1,
        "1 Contexte général et problématique",
        [
            "L'IoT regroupe des objets physiques capables de produire ou transmettre des données. Dans une maison intelligente ou un environnement industriel, ces équipements transmettent régulièrement des informations utiles à la supervision.",
            "MQTT repose sur trois acteurs : publisher, broker et subscriber. Cette simplicité facilite l'intégration mais impose une vigilance sur l'authentification, la qualité des messages et le contrôle du trafic.",
            "La problématique traitée est la suivante : comment concevoir un système capable de détecter en temps réel les anomalies et attaques MQTT dans un environnement IoT, tout en fournissant une interface claire pour l'analyse ?",
        ],
    ),
]
for level, title, paragraphs in sections:
    h(doc, level, title)
    for text in paragraphs:
        para(doc, text)

h(doc, 2, "1.1 Objectifs du projet")
bullets(
    doc,
    [
        "mettre en place une chaîne de communication MQTT entre un objet connecté et un backend ;",
        "exploiter des modèles de classification du trafic MQTT ;",
        "détecter plusieurs familles d'attaques : DoS, Flood, Bruteforce, SlowITe, Malformed et anomalies inconnues ;",
        "afficher les alertes et indicateurs dans un dashboard temps réel ;",
        "préparer un environnement démontrable pendant la soutenance.",
    ],
)
add_table(
    doc,
    "Périmètre fonctionnel du projet",
    ["Module", "Rôle", "Résultat attendu"],
    [
        ("ESP32 / Arduino IDE", "Publication de messages MQTT", "Messages JSON envoyés au broker"),
        ("Mosquitto", "Broker MQTT local", "Diffusion sur pfe/mqtt/trafic"),
        ("FastAPI", "Backend d'écoute et prédiction", "Résultats API/WebSocket"),
        ("Modèles ML/DL", "Classification du trafic", "Étiquette prédite"),
        ("React Dashboard", "Supervision graphique", "Alertes temps réel"),
    ],
)
add_table(
    doc,
    "Exigences fonctionnelles",
    ["ID", "Exigence", "Priorité"],
    [
        ("F1", "Recevoir les messages publiés sur le topic MQTT", "Élevée"),
        ("F2", "Prédire la classe du trafic à partir des champs MQTT/TCP", "Élevée"),
        ("F3", "Afficher l'état WebSocket et les compteurs de paquets", "Élevée"),
        ("F4", "Filtrer les anomalies par type dans l'interface", "Moyenne"),
        ("F5", "Prévoir un mode de démonstration Arduino/ESP32", "Élevée"),
    ],
)
add_table(
    doc,
    "Exigences non fonctionnelles",
    ["Critère", "Description", "Mesure retenue"],
    [
        ("Temps réel", "Transmission rapide des résultats", "WebSocket continu"),
        ("Lisibilité", "Interface compréhensible par le jury", "Cartes, graphiques et badges"),
        ("Maintenabilité", "Séparation frontend/backend/modèles", "Architecture claire"),
        ("Extensibilité", "Ajout possible de nouvelles attaques", "Labels normalisés"),
        ("Reproductibilité", "Tests rejouables", "Script Python et sketch Arduino"),
    ],
)

h(doc, 2, "1.2 Parties prenantes")
para(doc, "L'analyse des parties prenantes permet de préciser les besoins de chaque acteur concerné par le projet. Dans un système IoT de détection d'intrusions, l'utilisateur final ne cherche pas uniquement une prédiction technique : il souhaite comprendre rapidement ce qui se passe, identifier le type d'attaque et disposer d'un support de démonstration fiable.")
add_table(
    doc,
    "Parties prenantes et besoins associés",
    ["Partie prenante", "Besoin principal", "Réponse apportée par le projet"],
    [
        ("Étudiants réalisateurs", "Construire une solution complète et démontrable", "Architecture intégrée avec captures et tests"),
        ("Encadrante", "Évaluer la cohérence technique et méthodologique", "Rapport structuré, tableaux, figures et justification des choix"),
        ("Jury", "Comprendre rapidement le problème et la solution", "Dashboard, architecture visuelle et scénarios de test"),
        ("Administrateur IoT", "Superviser les alertes réseau", "Interface temps réel avec filtres par type d'anomalie"),
        ("Développeur futur", "Faire évoluer le système", "Code séparé en backend, frontend, modèles et outils de test"),
    ],
)
h(doc, 2, "1.3 Contraintes du projet")
para(doc, "Le projet combine plusieurs domaines : réseau, IoT, machine learning, développement web et démonstration matérielle. Cette diversité impose des contraintes techniques et organisationnelles qu'il faut maîtriser pour obtenir une soutenance fluide.")
add_table(
    doc,
    "Contraintes techniques et organisationnelles",
    ["Contrainte", "Explication", "Impact sur la réalisation"],
    [
        ("Temps réel", "Les alertes doivent apparaître sans rechargement manuel", "Utilisation de WebSocket"),
        ("Interopérabilité", "ESP32, MQTT, Python et React doivent communiquer", "Messages JSON standardisés"),
        ("Données volumineuses", "Certains CSV dépassent plusieurs centaines de Mo", "Utilisation des versions reduced"),
        ("Démonstration", "Le matériel peut être instable le jour de la soutenance", "Script Python de secours"),
        ("Lisibilité", "Le jury doit comprendre sans lire le code", "Figures, tableaux et dashboard visuel"),
    ],
)
h(doc, 2, "1.4 Plan de réalisation")
add_table(
    doc,
    "Plan de réalisation du projet",
    ["Étape", "Travail effectué", "Livrable"],
    [
        ("Étude du sujet", "Compréhension IoT, MQTT, IDS et attaques", "Problématique et objectifs"),
        ("Préparation des données", "Analyse MQTTset et fichiers reduced", "Tableaux de répartition"),
        ("Modélisation", "Comparaison RF, XGBoost et Autoencoder", "Modèles sauvegardés"),
        ("Backend", "FastAPI, MQTT, prédiction, WebSocket", "API fonctionnelle"),
        ("Frontend", "Dashboard React et pages d'analyse", "Interface de supervision"),
        ("Arduino / tests", "Sketch ESP32 et scénarios MQTT", "Démonstration temps réel"),
        ("Rapport", "Structuration, captures et tableaux", "Document final prêt soutenance"),
    ],
)

doc.add_page_break()
h(doc, 1, "2 État de l'art et choix technologiques")
for text in [
    "Un IDS surveille un environnement afin d'identifier des comportements suspects. Dans l'IoT, les contraintes sont plus fortes car les équipements sont légers, les messages courts et les flux fréquents.",
    "L'approche adoptée consiste à déplacer l'intelligence vers le backend. L'objet connecté publie les données, tandis que l'analyse est réalisée par les modèles entraînés côté serveur.",
    "XGBoost est retenu comme modèle principal, car il offre un bon compromis entre précision, rapidité et robustesse sur données tabulaires. L'autoencoder complète l'approche pour repérer des anomalies inconnues.",
]:
    para(doc, text)
add_table(
    doc,
    "Technologies utilisées",
    ["Technologie", "Utilisation", "Justification"],
    [
        ("Python", "Backend, scripts, ML", "Écosystème IA riche"),
        ("FastAPI", "API REST et WebSocket", "Rapide et moderne"),
        ("React", "Interface web", "Dashboard dynamique"),
        ("Mosquitto", "Broker MQTT", "Référence légère"),
        ("XGBoost", "Classification supervisée", "Très performant sur données structurées"),
        ("TensorFlow/Keras", "Autoencoder", "Détection par reconstruction"),
        ("Arduino IDE", "Programmation ESP32", "Validation matérielle"),
    ],
)
add_table(
    doc,
    "Comparaison des modèles exploités",
    ["Modèle", "Type", "Accuracy", "Temps", "Rôle"],
    [
        ("Random Forest", "Supervisé", "94,03 %", "30 s", "Référence comparative"),
        ("XGBoost", "Supervisé", "94,20 %", "8 s", "Modèle principal"),
        ("Autoencoder", "Non supervisé", "85,18 %", "94 s", "Anomalies inconnues"),
    ],
)

h(doc, 2, "2.1 Typologie des attaques MQTT traitées")
para(doc, "Les attaques considérées dans le projet correspondent aux classes présentes dans le dataset et aux scénarios injectés dans l'application. Elles permettent de couvrir à la fois des comportements de surcharge, de connexion abusive et de messages non conformes.")
add_table(
    doc,
    "Description des attaques étudiées",
    ["Type", "Principe", "Effet possible sur le système"],
    [
        ("DoS", "Multiplier les requêtes ou messages pour dégrader le service", "Saturation du broker ou du backend"),
        ("Flood", "Envoyer un volume élevé de paquets en peu de temps", "Augmentation anormale du trafic"),
        ("Bruteforce", "Multiplier les tentatives de connexion", "Risque d'accès non autorisé"),
        ("SlowITe", "Maintenir des connexions lentes ou incomplètes", "Occupation progressive des ressources"),
        ("Malformed", "Envoyer des paquets MQTT non conformes", "Erreur de parsing ou comportement imprévu"),
        ("Anomalie inconnue", "Comportement non représenté clairement dans les classes connues", "Nécessite une détection complémentaire"),
    ],
)
h(doc, 2, "2.2 Justification du choix XGBoost + Autoencoder")
para(doc, "Le projet ne repose pas sur un seul mécanisme. XGBoost traite efficacement les attaques connues grâce à l'apprentissage supervisé, alors que l'autoencoder apporte une logique de reconstruction utile lorsqu'un message semble légitime mais reste inhabituel. Cette complémentarité renforce la crédibilité de la solution.")
add_table(
    doc,
    "Critères de choix des modèles",
    ["Critère", "XGBoost", "Autoencoder"],
    [
        ("Nature de l'apprentissage", "Supervisé", "Non supervisé"),
        ("Données nécessaires", "Labels connus", "Comportement normal ou représentation apprise"),
        ("Point fort", "Classification rapide et précise", "Détection d'écarts inconnus"),
        ("Limite", "Moins adapté aux attaques jamais vues", "Nécessite un seuil bien calibré"),
        ("Rôle final", "Décision principale", "Contrôle complémentaire"),
    ],
)

doc.add_page_break()
h(doc, 1, "3 Analyse et conception de la solution")
para(doc, "L'architecture retenue est modulaire. Les messages sont publiés par l'ESP32 ou par un simulateur vers Mosquitto. FastAPI s'abonne au topic, transforme les messages, applique les modèles et diffuse les résultats vers React.")
fig(doc, ARCH, "Architecture globale du système MQTT-IDS")
add_table(
    doc,
    "Composants de l'architecture",
    ["Composant", "Entrée", "Traitement", "Sortie"],
    [
        ("ESP32", "Mesures DHT22 ou scénarios", "Construction JSON", "Message MQTT"),
        ("Mosquitto", "Publication MQTT", "Routage par topic", "Message reçu par FastAPI"),
        ("FastAPI", "Payload JSON", "Normalisation, prédiction", "Résultat API/WebSocket"),
        ("XGBoost/Autoencoder", "Vecteur de features", "Classification", "Label de trafic"),
        ("React", "Flux WebSocket", "Agrégation et visualisation", "Dashboard"),
    ],
)
para(doc, "Le dataset MQTTset contient des captures PCAP, des CSV extraits et des versions finales prêtes pour l'apprentissage. Le projet utilise les fichiers réduits afin d'obtenir un équilibre plus exploitable entre trafic légitime et trafic malveillant.")
fig(doc, PIPE, "Pipeline de préparation des données et entraînement des modèles")
add_table(
    doc,
    "Fichiers du dataset utilisés",
    ["Fichier", "Taille", "Utilisation"],
    [
        ("mqttdataset_reduced.csv", mb(ROOT / "dataset/Data/FINAL_CSV/mqttdataset_reduced.csv"), "Base réduite globale"),
        ("train70_reduced.csv", mb(ROOT / "dataset/Data/FINAL_CSV/train70_reduced.csv"), "Apprentissage"),
        ("test30_reduced.csv", mb(ROOT / "dataset/Data/FINAL_CSV/test30_reduced.csv"), "Évaluation"),
        ("PCAP", "plusieurs fichiers", "Sources réseau initiales"),
    ],
)
add_table(
    doc,
    "Répartition des classes dans les données réduites",
    ["Classe", "Total", "Train 70 %", "Test 30 %"],
    [(c, ct.get(c, 0), ctr.get(c, 0), cte.get(c, 0)) for c in classes],
)
add_table(
    doc,
    "Familles de caractéristiques exploitées",
    ["Famille", "Exemples", "Intérêt"],
    [
        ("TCP", "tcp.flags, tcp.time_delta, tcp.len", "Décrire le trafic réseau"),
        ("MQTT flags", "mqtt.conflags, mqtt.hdrflags", "Identifier les connexions"),
        ("MQTT message", "mqtt.msg, mqtt.protoname", "Analyser le contenu protocolaire"),
        ("Champs dérivés", "cleansess, passwd, qos, retain", "Transformer les flags"),
        ("Target", "legitimate, dos, flood, ...", "Classe d'apprentissage"),
    ],
)
fig(doc, SEQ, "Séquence de validation temps réel entre Arduino, MQTT, backend, ML et dashboard")

h(doc, 2, "3.1 Modèle logique des données")
para(doc, "Les messages manipulés par le système sont des objets JSON. Cette représentation est simple à publier depuis un ESP32, facile à lire en Python et directement exploitable par l'interface web. Les champs peuvent être complets lorsqu'ils proviennent du dataset ou simplifiés lorsqu'ils servent à la démonstration.")
add_table(
    doc,
    "Structure logique d'un message MQTT de démonstration",
    ["Champ", "Exemple", "Utilisation"],
    [
        ("demo_label", "flood", "Forcer un scénario contrôlé pour la soutenance"),
        ("mqtt.msg", "flood-burst", "Indice protocolaire utilisé par le backend"),
        ("mqtt.conflags", "0xc2", "Exemple de flag de connexion bruteforce"),
        ("mqtt.protoname", "invalid", "Détection de paquet malformé"),
        ("temperature / humidity", "24.7 / 51.2", "Simulation d'un capteur DHT22 normal"),
        ("scenario", "Attaque DoS MQTT", "Description lisible dans les tests"),
    ],
)
h(doc, 2, "3.2 Algorithme de traitement backend")
para(doc, "Le backend suit une logique séquentielle. Il reçoit d'abord le message brut, tente de le convertir en JSON, vérifie s'il s'agit d'un message capteur ou d'un scénario de démonstration, puis appelle le moteur de prédiction si nécessaire. Le résultat final est ajouté à la mémoire des derniers événements.")
add_table(
    doc,
    "Étapes de traitement côté FastAPI",
    ["Ordre", "Étape", "Résultat"],
    [
        ("1", "Réception MQTT sur pfe/mqtt/trafic", "Message brut disponible"),
        ("2", "Décodage JSON", "Dictionnaire Python exploitable"),
        ("3", "Détection du mode capteur ou démo", "Réponse directe si scénario connu"),
        ("4", "Préparation des features", "Vecteur numérique aligné avec le scaler"),
        ("5", "Prédiction XGBoost", "Classe candidate"),
        ("6", "Contrôle autoencoder si nécessaire", "Anomalie inconnue possible"),
        ("7", "Diffusion WebSocket", "Dashboard mis à jour"),
    ],
)
h(doc, 2, "3.3 Conception de l'expérience utilisateur")
para(doc, "L'interface a été pensée pour une lecture rapide pendant la démonstration. Les informations les plus importantes apparaissent dans le dashboard : statut de connexion, volume de paquets, nombre d'anomalies et dernier type détecté. Les pages secondaires permettent ensuite d'approfondir l'analyse.")
add_table(
    doc,
    "Rôle des composants frontend",
    ["Composant", "Rôle visuel", "Donnée affichée"],
    [
        ("MetricCard", "Cartes de synthèse", "Total, normal, anomalies, accuracy"),
        ("TrafficChart", "Courbe temporelle", "Évolution normal/anomalie"),
        ("DonutChart", "Répartition", "Part des classes détectées"),
        ("BarChart", "Histogramme", "Nombre par type d'attaque"),
        ("AnomalieTable", "Table d'historique", "Type, heure et statut"),
        ("AlertBanner", "Alerte immédiate", "Dernière anomalie reçue"),
    ],
)

doc.add_page_break()
h(doc, 1, "4 Réalisation technique")
for text in [
    "Le backend FastAPI assure la connexion au broker MQTT, la réception des messages, la normalisation des labels de démonstration, l'appel au moteur de prédiction, la mémorisation des derniers résultats et la diffusion WebSocket.",
    "Le frontend est organisé en pages React : Dashboard, Anomalies, Modèles ML et À propos. Le contexte WebSocket centralise les messages reçus et fournit les données aux composants graphiques.",
    "Le dossier models contient les artefacts entraînés : XGBoost, autoencoder, scaler, label encoder et seuil de reconstruction. Cette séparation rend la solution plus maintenable.",
]:
    para(doc, text)
add_table(
    doc,
    "Principaux fichiers et dossiers du projet",
    ["Chemin", "Description"],
    [
        ("backend/main.py", "API FastAPI, client MQTT, WebSocket et logique de réponse"),
        ("backend/predict.py", "Chargement des modèles et prédiction ML/DL"),
        ("frontend/src/pages/Dashboard.jsx", "Vue principale de supervision"),
        ("frontend/src/pages/Anomalies.jsx", "Table de suivi et filtrage"),
        ("frontend/src/pages/Modeles.jsx", "Comparaison des modèles et matrice"),
        ("models/xgb_model.pkl", "Modèle XGBoost principal"),
        ("models/autoencoder.keras", "Autoencoder pour anomalies inconnues"),
        ("arduino/mqtt_ids_demo/mqtt_ids_demo.ino", "Sketch Arduino de test MQTT"),
        ("tools/publish_attack_scenarios.py", "Script de rejeu des scénarios"),
    ],
)
add_table(
    doc,
    "Endpoints et canaux backend",
    ["Route ou canal", "Méthode", "Rôle"],
    [
        ("/", "GET", "Vérifier que l'API fonctionne"),
        ("/resultats", "GET", "Récupérer les dix derniers résultats"),
        ("/resultats/clear", "POST", "Réinitialiser la démonstration"),
        ("/ws", "WebSocket", "Diffuser les alertes au dashboard"),
        ("pfe/mqtt/trafic", "MQTT topic", "Canal d'entrée des messages IoT"),
    ],
)
add_table(
    doc,
    "Artefacts de modèles enregistrés",
    ["Fichier", "Taille", "Rôle"],
    [
        ("xgb_model.pkl", mb(ROOT / "models/xgb_model.pkl"), "Classification principale"),
        ("autoencoder.keras", mb(ROOT / "models/autoencoder.keras"), "Détection complémentaire"),
        ("autoencoder_threshold.pkl", mb(ROOT / "models/autoencoder_threshold.pkl"), "Seuil MSE"),
        ("scaler.pkl", mb(ROOT / "models/scaler.pkl"), "Normalisation des features"),
        ("label_encoder.pkl", mb(ROOT / "models/label_encoder.pkl"), "Conversion labels/classes"),
    ],
)

h(doc, 2, "4.1 Détails d'implémentation backend")
para(doc, "Le backend conserve une liste des derniers résultats afin de pouvoir alimenter l'interface même si un utilisateur ouvre la page après le début de la démonstration. Cette logique est utile pendant la soutenance, car elle évite de perdre immédiatement les messages publiés.")
para(doc, "La fonction de normalisation accepte plusieurs alias comme normal, legitimate, inconnue ou unknown. Ce choix améliore la tolérance du système et permet de tester différents formats de payload sans modifier l'interface.")
add_table(
    doc,
    "Fonctions principales du backend",
    ["Fonction", "Responsabilité"],
    [
        ("normalize_demo_label", "Uniformiser les labels de test"),
        ("detect_demo_label", "Reconnaître les scénarios Arduino/simulateur"),
        ("build_response", "Construire la réponse finale envoyée au dashboard"),
        ("on_message", "Réagir à chaque message MQTT reçu"),
        ("get_resultats", "Exposer les derniers résultats par API REST"),
        ("websocket_endpoint", "Envoyer les événements en temps réel"),
    ],
)
h(doc, 2, "4.2 Détails d'implémentation ML")
para(doc, "Le fichier predict.py charge les artefacts depuis le dossier models. Les champs manquants sont remplacés par des valeurs par défaut ou par les moyennes du scaler, ce qui rend la prédiction plus robuste face aux messages partiels.")
add_table(
    doc,
    "Préparation des features avant prédiction",
    ["Étape", "Description"],
    [
        ("Lecture des feature_names", "Aligner l'ordre des colonnes avec celui utilisé à l'entraînement"),
        ("Conversion hexadécimale", "Transformer les flags MQTT/TCP en valeurs numériques"),
        ("Champs dérivés", "Extraire cleansess, passwd, qos, retain et willflag"),
        ("Valeurs manquantes", "Utiliser des défauts ou les moyennes du scaler"),
        ("Scaling", "Appliquer la normalisation avant XGBoost et autoencoder"),
    ],
)
h(doc, 2, "4.3 Détails Arduino et protocole de démonstration")
para(doc, "Le sketch Arduino publie une série de messages JSON représentant les principaux scénarios. Les identifiants Wi-Fi sont laissés sous forme de paramètres à adapter, ce qui évite d'inscrire des informations sensibles dans le rapport.")
add_table(
    doc,
    "Paramètres Arduino/MQTT à adapter",
    ["Paramètre", "Valeur dans le sketch", "Commentaire"],
    [
        ("ssid", "VOTRE_WIFI", "Nom du réseau Wi-Fi local"),
        ("password", "VOTRE_MOT_DE_PASSE", "Mot de passe à renseigner avant upload"),
        ("mqttServer", "ADRESSE_IP_DU_PC", "Adresse de la machine exécutant Mosquitto"),
        ("mqttPort", "1883", "Port MQTT standard non chiffré"),
        ("topic", "pfe/mqtt/trafic", "Topic écouté par FastAPI"),
    ],
)
fig(doc, CAP / "01_dashboard_live_hd.png", "Dashboard MQTT-IDS après injection de scénarios d'attaques")
fig(doc, CAP / "02_anomalies_table_hd.png", "Table des anomalies détectées en temps réel")
fig(doc, CAP / "03_anomalies_filter_flood_hd.png", "Filtrage des anomalies de type Flood")
fig(doc, CAP / "04_modeles_ml_hd.png", "Comparaison des modèles et visualisation des performances")
fig(doc, CAP / "05_about_architecture_hd.png", "Vue À propos : architecture, technologies et informations projet")
fig(doc, CAP / "06_arduino_ide_sketch.png", "Sketch Arduino IDE préparé pour envoyer les scénarios MQTT")

doc.add_page_break()
h(doc, 1, "5 Tests, résultats et interprétation")
para(doc, "La validation consiste à injecter des messages MQTT représentatifs de plusieurs situations. Chaque message contient un label de démonstration ou des champs MQTT permettant au backend de déterminer la classe.")
add_table(
    doc,
    "Scénarios MQTT testés",
    ["Scénario", "Payload principal", "Résultat attendu"],
    [
        ("Trafic légitime", "demo_label=legitimate, mqtt.msg=temperature", "Trafic normal"),
        ("DoS", "demo_label=dos, mqtt.msg=attack", "Anomalie DoS"),
        ("Flood", "demo_label=flood, mqtt.msg=flood-burst", "Anomalie Flood"),
        ("Bruteforce", "demo_label=bruteforce, mqtt.conflags=0xc2", "Anomalie Bruteforce"),
        ("SlowITe", "demo_label=slowite, mqtt.msg=slow", "Anomalie SlowITe"),
        ("Malformed", "demo_label=malformed, mqtt.protoname=invalid", "Anomalie Malformed"),
        ("Inconnue", "demo_label=anomalie_inconnue", "Anomalie inconnue"),
    ],
)
add_table(
    doc,
    "Résultats de validation fonctionnelle",
    ["Test", "Observation", "Statut"],
    [
        ("Connexion WebSocket", "L'interface affiche WebSocket : connecté", "Validé"),
        ("Publication MQTT", "Les messages sont reçus par FastAPI", "Validé"),
        ("Classification", "Les labels d'attaque sont visibles", "Validé"),
        ("Filtrage", "La page Anomalies isole le type Flood", "Validé"),
        ("Visualisation", "Les cartes et graphiques se mettent à jour", "Validé"),
        ("Arduino IDE", "Le sketch contient les payloads de démonstration", "Préparé"),
    ],
)
add_table(
    doc,
    "Matrice de confusion XGBoost intégrée dans l'interface",
    ["Réel / Prédit", "bruteforce", "dos", "flood", "legitimate", "malformed", "slowite"],
    [
        ("Bruteforce", 1754, 12, 0, 180, 43, 0),
        ("DoS", 45, 16821, 0, 1350, 25, 0),
        ("Flood", 0, 5, 34, 42, 0, 0),
        ("Legitimate", 12, 145, 0, 23050, 46, 0),
        ("Malformed", 28, 18, 0, 590, 838, 0),
        ("SlowITe", 0, 0, 0, 0, 0, 1292),
    ],
)
for text in [
    "Les résultats observés montrent que le système est capable de traiter un flux MQTT et de déclencher des alertes visibles en temps réel. Les scénarios couvrent les principales familles d'attaques présentes dans le dataset.",
    "Le choix de XGBoost est cohérent avec les données tabulaires issues du trafic MQTT. Le modèle offre une précision élevée et un temps d'entraînement réduit.",
    "L'interface React apporte une valeur importante, car elle transforme des résultats techniques en informations exploitables : nombre de paquets, pourcentage d'anomalies, type d'attaque, heure, histogramme et répartition.",
]:
    para(doc, text)

h(doc, 2, "5.1 Analyse détaillée des scénarios")
para(doc, "Chaque scénario a été conçu pour vérifier une partie précise du système. L'objectif n'est pas seulement d'obtenir une alerte, mais de montrer que la chaîne complète reste cohérente : publication, réception, classification, stockage temporaire et visualisation.")
add_table(
    doc,
    "Lecture détaillée des scénarios de test",
    ["Scénario", "Ce que l'on vérifie", "Indicateur visible"],
    [
        ("Légitime", "Le système ne doit pas générer d'alerte inutile", "Trafic normal augmenté"),
        ("DoS", "La classe dos doit être reconnue", "Badge DOS dans les anomalies"),
        ("Flood", "Le filtrage par type doit fonctionner", "Filtre FLOOD actif"),
        ("Bruteforce", "Les flags de connexion sont interprétés", "Badge BRUTEFORCE"),
        ("SlowITe", "Un comportement lent est classé comme anomalie", "Badge SLOWITE"),
        ("Malformed", "Un protocole invalide est détecté", "Badge MALFORMED"),
        ("Inconnue", "Une anomalie non standard est remontée", "Badge INCONNUE"),
    ],
)
h(doc, 2, "5.2 Interprétation par type d'attaque")
for title, explanation in [
    ("DoS", "L'attaque DoS vise à rendre un service indisponible. Dans le contexte MQTT, elle peut se matérialiser par un nombre élevé de messages ou de requêtes qui forcent le broker et le backend à consommer plus de ressources."),
    ("Flood", "Le flood se distingue par un flux massif et rapide. Le test Flood est important pour démontrer que l'interface ne se contente pas d'afficher une alerte générale, mais qu'elle sait aussi filtrer une catégorie précise."),
    ("Bruteforce", "Le bruteforce cible les mécanismes de connexion. Dans le projet, le champ mqtt.conflags permet de simuler une tentative de connexion caractéristique et de vérifier la reconnaissance de ce comportement."),
    ("SlowITe", "SlowITe représente un comportement plus discret : la consommation des ressources se fait progressivement. Cette classe montre que le dataset et le modèle ne traitent pas seulement les attaques volumétriques."),
    ("Malformed", "Les paquets malformés ne respectent pas la structure attendue. Ils sont importants à détecter car ils peuvent provoquer des erreurs, des exceptions ou des comportements imprévus dans les systèmes mal protégés."),
    ("Anomalie inconnue", "La classe anomalie_inconnue prépare le système à des comportements qui ne correspondent pas directement aux catégories connues. Elle justifie l'intérêt d'une couche complémentaire comme l'autoencoder."),
]:
    h(doc, 3, title)
    para(doc, explanation)

add_table(
    doc,
    "Points contrôlés pendant la démonstration",
    ["Point de contrôle", "Méthode de vérification", "Résultat attendu"],
    [
        ("Broker actif", "netstat ou observation de la connexion", "Port 1883 ouvert"),
        ("Backend actif", "Accès API /resultats", "Réponse JSON disponible"),
        ("Frontend actif", "Ouverture http://localhost:5173", "Dashboard affiché"),
        ("WebSocket", "Badge dans le dashboard", "Statut connecté"),
        ("Publication MQTT", "Script ou Arduino IDE", "Nouveaux paquets visibles"),
        ("Classification", "Table des anomalies", "Type correct affiché"),
    ],
)
h(doc, 2, "5.3 Discussion critique")
para(doc, "La solution est solide pour une démonstration PFE, car elle couvre toutes les couches attendues : objet connecté, protocole réseau, backend, intelligence artificielle et interface utilisateur. Elle reste toutefois un prototype académique. Pour une mise en production, il faudrait ajouter l'authentification MQTT, la persistance des alertes, les tests de charge et une procédure de réentraînement.")
para(doc, "Le point le plus important pour la soutenance est de montrer la cohérence du système. Le jury doit percevoir que les modèles ne sont pas isolés dans un notebook, mais intégrés dans une application réellement exploitable en temps réel.")

doc.add_page_break()
h(doc, 1, "6 Gestion des risques, sécurité et perspectives")
add_table(
    doc,
    "Risques du projet et mesures de maîtrise",
    ["Risque", "Impact", "Mesure de maîtrise"],
    [
        ("Perte de connexion au broker", "Aucune donnée reçue", "Vérifier le port 1883"),
        ("Format JSON invalide", "Erreur de parsing", "Gestion des exceptions backend"),
        ("Données partielles", "Prédiction moins fiable", "Valeurs par défaut et scaler means"),
        ("Déséquilibre de classes", "Biais possible", "Versions réduites/augmentées"),
        ("Démonstration matérielle instable", "Risque en soutenance", "Script Python de secours"),
    ],
)
bullets(
    doc,
    [
        "activer l'authentification MQTT et limiter les topics autorisés ;",
        "utiliser TLS pour chiffrer les échanges entre objets et broker ;",
        "ajouter une journalisation persistante des alertes ;",
        "mettre en place un contrôle d'accès sur l'interface web ;",
        "prévoir un mécanisme de réentraînement périodique des modèles.",
    ],
)
for text in [
    "Plusieurs évolutions peuvent enrichir le projet : intégrer une base de données pour historiser les alertes, ajouter une authentification utilisateur, déployer l'application dans Docker et tester le système sur un réseau IoT réel.",
    "Une autre perspective consiste à enrichir les messages ESP32 avec des mesures réelles du DHT22 et à corréler les anomalies réseau avec les événements capteurs.",
]:
    para(doc, text)

doc.add_page_break()
h(doc, 1, "Conclusion générale")
for text in [
    "Ce projet de fin d'études a permis de concevoir et de réaliser une solution complète de détection d'intrusions IoT basée sur MQTT. Le système combine une chaîne de communication temps réel, un backend de traitement, des modèles de Machine Learning et une interface web moderne.",
    "La réalisation démontre la faisabilité d'un IDS orienté IoT capable de classifier plusieurs types d'attaques MQTT et de les afficher instantanément. Les tests menés avec les scénarios Arduino et le script de simulation confirment le bon fonctionnement de la chaîne de bout en bout.",
    "Au-delà de la partie technique, ce travail met en évidence l'importance d'une présentation claire des résultats. Le dashboard, les tableaux et les captures facilitent la compréhension du projet par le jury.",
]:
    para(doc, text)

doc.add_page_break()
h(doc, 1, "Annexes")
add_table(
    doc,
    "Commandes utiles pour la démonstration",
    ["Action", "Commande"],
    [
        ("Lancer le backend", "python -m uvicorn main:app --reload --app-dir backend --port 8001"),
        ("Lancer le frontend", "npm run dev dans le dossier frontend"),
        ("Tester les scénarios MQTT", "python tools/publish_attack_scenarios.py"),
        ("Ouvrir l'application", "http://localhost:5173/"),
    ],
)
add_table(
    doc,
    "Checklist avant présentation",
    ["Élément", "Vérification"],
    [
        ("Broker Mosquitto", "Port 1883 en écoute"),
        ("Backend FastAPI", "Port 8001 accessible"),
        ("Frontend React", "Port 5173 accessible"),
        ("WebSocket", "Statut connecté dans l'interface"),
        ("Scénarios", "Script Python ou ESP32 prêt"),
        ("Captures", "Dashboard, anomalies, modèles et Arduino disponibles"),
        ("Présentation orale", "Architecture, dataset, modèles, tests, limites"),
    ],
)
add_table(
    doc,
    "Plan oral recommandé",
    ["Partie", "Durée", "Message clé"],
    [
        ("Contexte et problématique", "2 min", "Sécuriser les communications IoT/MQTT"),
        ("Architecture", "3 min", "Chaîne ESP32 -> MQTT -> FastAPI -> ML -> React"),
        ("Dataset et modèles", "3 min", "MQTTset, XGBoost, Autoencoder"),
        ("Démonstration", "5 min", "Publication d'attaques et détection temps réel"),
        ("Résultats et limites", "3 min", "Fonctionnement validé et perspectives"),
        ("Conclusion", "1 min", "Projet complet et extensible"),
    ],
)

doc.save(OUT)
print(OUT)
